#!/usr/bin/env python3
"""Download ITU-T X.680-683 (basic notation) and X.690/691/692/693/696/697
(encoding rules) specs and extract EBNF grammar / lexical tokens.

Steps performed automatically:
  1. Download PDFs and Word (.doc) files from itu.int (skipped if already present).
  2. Convert .doc → .docx via LibreOffice (skipped if .docx already present).
  3. Extract grammar (BNF/BNF-Continue styles) → asn1.ebnf
  4. Extract lexical tokens                    → asn1.tokens
  5. Convert each spec to Markdown             → <name>.md (e.g. X6801.md)

Prerequisites:
    pip install python-docx
    libreoffice   (for .doc → .docx conversion)

Usage:
    python3 extract_grammar.py                   # download + convert + extract
    python3 extract_grammar.py --clean           # remove downloaded artifacts
    python3 extract_grammar.py [file.docx ...]   # explicit files, skip download

Output: asn1.ebnf, asn1.tokens, and one <name>.md per .docx, all written next
to this script.
"""

import os
import re
import subprocess
import sys
import urllib.error
import urllib.request

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

SPECS_DIR = os.path.dirname(os.path.abspath(__file__))
GRAMMAR_FILE = os.path.join(SPECS_DIR, "asn1.ebnf")
TOKENS_FILE  = os.path.join(SPECS_DIR, "asn1.tokens")

# ── download / convert ────────────────────────────────────────────────────────

_BASE_URL = "https://www.itu.int/rec/dologin_pub.asp"

_DOWNLOADS = [
    # (rec_id,          pdf_filename,       doc_filename,  docx_filename)
    # -- basic notation / information objects / constraints / parameterization --
    ("T-REC-X.680-202102-I", "X.680-202102.pdf", "X6801.doc", "X6801.docx"),
    ("T-REC-X.681-202102-I", "X.681-202102.pdf", "X6811.doc", "X6811.docx"),
    ("T-REC-X.682-202102-I", "X.682-202102.pdf", "X6821.doc", "X6821.docx"),
    ("T-REC-X.683-202102-I", "X.683-202102.pdf", "X6831.doc", "X6831.docx"),
    # -- encoding rules: BER/CER/DER, PER, ECN, XER, OER, JER --
    # ITU does not publish a .doc source for this series (dologin_pub.asp
    # returns a genuine HTTP 500 for the DOC format, confirmed by hand), so
    # doc/docx are None: only the PDF is fetched, and grammar/token/markdown
    # extraction (which all require .docx) are skipped for these.
    ("T-REC-X.690-202102-I", "X.690-202102.pdf", None, None),
    ("T-REC-X.691-202102-I", "X.691-202102.pdf", None, None),
    ("T-REC-X.692-202102-I", "X.692-202102.pdf", None, None),
    ("T-REC-X.693-202102-I", "X.693-202102.pdf", None, None),
    ("T-REC-X.696-202102-I", "X.696-202102.pdf", None, None),
    ("T-REC-X.697-202102-I", "X.697-202102.pdf", None, None),
]

_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (X11; Linux x86_64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": "application/octet-stream,*/*",
    "Referer": "https://www.itu.int/",
}


def _fetch(url: str, dest: str) -> bool:
    """Download *url* to *dest*. Returns True on success."""
    req = urllib.request.Request(url, headers=_HEADERS)
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = resp.read()
        with open(dest, "wb") as f:
            f.write(data)
        print(f"  [ok] {os.path.basename(dest)} ({len(data) // 1024} KB)")
        return True
    except urllib.error.HTTPError as exc:
        print(f"  [error] HTTP {exc.code} fetching {url}", file=sys.stderr)
        return False
    except urllib.error.URLError as exc:
        print(f"  [error] {exc.reason} fetching {url}", file=sys.stderr)
        return False


def _rec_url(rec_id: str, fmt: str) -> str:
    return f"{_BASE_URL}?lang=e&id={rec_id}!!{fmt}-E&type=items"


def download_all() -> list[str]:
    """Download PDFs + .doc files and convert to .docx. Returns list of .docx paths."""
    docx_paths: list[str] = []
    ok = True

    for rec_id, pdf_name, doc_name, docx_name in _DOWNLOADS:
        pdf_path = os.path.join(SPECS_DIR, pdf_name)

        # 1. PDF
        if os.path.exists(pdf_path):
            print(f"[skip] {pdf_name} already exists")
        else:
            print(f"[download] {pdf_name}")
            ok = _fetch(_rec_url(rec_id, "PDF"), pdf_path) and ok

        # No .doc source published for this recommendation: nothing further
        # to download or convert (grammar/token/markdown extraction, which
        # all require .docx, are simply unavailable for it).
        if doc_name is None:
            print(f"[skip] {pdf_name}: no .doc source published, PDF only")
            continue

        doc_path  = os.path.join(SPECS_DIR, doc_name)
        docx_path = os.path.join(SPECS_DIR, docx_name)

        # 2. Word .doc
        if os.path.exists(doc_path):
            print(f"[skip] {doc_name} already exists")
        else:
            print(f"[download] {doc_name}")
            ok = _fetch(_rec_url(rec_id, "DOC"), doc_path) and ok

        # 3. Convert .doc → .docx
        if os.path.exists(docx_path):
            print(f"[skip] {docx_name} already exists")
        elif os.path.exists(doc_path):
            print(f"[convert] {doc_name} → {docx_name}")
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", doc_path],
                cwd=SPECS_DIR,
                capture_output=True,
                text=True,
            )
            if result.returncode != 0 or not os.path.exists(docx_path):
                print(f"  [error] libreoffice conversion failed:\n{result.stderr}", file=sys.stderr)
                ok = False
            else:
                print(f"  [ok] {docx_name}")
        else:
            print(f"[skip] {docx_name}: source .doc not available", file=sys.stderr)

        if os.path.exists(docx_path):
            docx_paths.append(docx_path)

    if not ok:
        print("Warning: some downloads or conversions failed.", file=sys.stderr)

    return sorted(docx_paths)


def clean() -> None:
    """Remove downloaded PDFs, .doc and .docx artifacts from the Specs directory."""
    for _rec_id, pdf_name, doc_name, docx_name in _DOWNLOADS:
        for name in (pdf_name, doc_name, docx_name):
            if name is None:
                continue
            path = os.path.join(SPECS_DIR, name)
            if os.path.exists(path):
                os.remove(path)
                print(f"[removed] {name}")
            else:
                print(f"[skip] {name} not present")


# ── helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()

def _table_cells(tbl_element) -> list[str]:
    cells = []
    for tc in tbl_element.iter(qn("w:tc")):
        val = "".join(n.text or "" for n in tc.iter(qn("w:t"))).strip()
        if val:
            cells.append(val)
    return cells

# ── grammar extraction ────────────────────────────────────────────────────────

_PROD_RE   = re.compile(r"^([A-Z][A-Za-z][A-Za-z0-9-]*)\s*::=(.*)", re.DOTALL)
_BNF_STYLES = {"BNF", "BNF Continue"}


def extract_grammar(docx_path: str) -> list[dict]:
    """Return unique grammar productions from a .docx file in first-appearance order."""
    doc = Document(docx_path)
    seen: set[str] = set()
    results: list[dict] = []
    current: dict | None = None

    def _commit():
        nonlocal current
        if current and current["name"] not in seen:
            seen.add(current["name"])
            results.append(current)
        current = None

    for para in doc.paragraphs:
        sname = para.style.name if para.style else ""
        if sname not in _BNF_STYLES:
            continue
        if not para.text.strip():
            continue

        if sname == "BNF":
            raw_lines = para.text.split("\n")
            first = _clean(raw_lines[0])
            m = _PROD_RE.match(first)
            if m:
                _commit()
                name = m.group(1)
                rest = _clean(m.group(2))
                head = f"{name} ::= {rest}" if rest else f"{name} ::="
                current = {"name": name, "lines": [head]}
                for raw in raw_lines[1:]:
                    line = _clean(raw)
                    if line:
                        current["lines"].append(
                            "| " + line[1:].strip() if line.startswith("|") else line
                        )
            else:
                _commit()
        else:  # BNF Continue
            if current is None:
                continue
            text = _clean(para.text)
            line = ("| " + text[1:].strip()) if text.startswith("|") else text
            current["lines"].append(line)

    _commit()
    return results


def _format_prod(prod: dict) -> str:
    lines = prod["lines"]
    if not lines:
        return ""
    out = [lines[0]]
    for line in lines[1:]:
        out.append(f"    {line}")
    return "\n".join(out)

# ── token extraction ──────────────────────────────────────────────────────────

# Matches: "Name of lexical item – foo"  or  "Name of item – "::=""
_TOKEN_HDR_RE = re.compile(r"^Name of (?:lexical item|item)\s*[–-]\s*(.+)$")

def extract_tokens(docx_path: str) -> dict:
    """
    Return a dict with keys:
      named    : list of {"name": str, "desc": str}
      single   : list of str  (e.g. '"{"')
      keywords : list of str  (reserved words)
    """
    doc = Document(docx_path)
    seen_names: set[str] = set()
    seen_single: set[str] = set()
    named: list[dict] = []
    single: list[str] = []
    keywords: list[str] = []

    # We need access to both paragraphs and tables in document order.
    # Iterate over body children directly.
    body = doc.element.body
    pending_token: dict | None = None   # token whose desc we are collecting
    in_single_char_block = False        # True between "Names of lexical items –" and its desc
    in_keywords_block = False           # True right after "Names of reserved words –"

    def _flush_token():
        nonlocal pending_token
        if pending_token and pending_token["name"] not in seen_names:
            seen_names.add(pending_token["name"])
            named.append(pending_token)
        pending_token = None

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "tbl":
            if in_keywords_block:
                keywords.extend(_table_cells(child))
                in_keywords_block = False
            continue

        if tag != "p":
            continue

        # Reconstruct paragraph text and style from the XML element
        style_node = child.find(".//" + qn("w:pStyle"))
        sname = style_node.get(qn("w:val"), "") if style_node is not None else ""
        raw_text = ""
        for n in child.iter():
            if n.tag == qn("w:t"):
                raw_text += n.text or ""
            elif n.tag == qn("w:tab"):
                raw_text += "\t"
        text = _clean(raw_text)

        if not text:
            continue

        # ── single-char block ──────────────────────────────────────────────
        if text == "Names of lexical items –" or text == "Names of lexical items -":
            in_single_char_block = True
            continue

        if in_single_char_block:
            if text.startswith("A lexical item with any"):
                in_single_char_block = False
                continue
            # Characters are concatenated in the XML (no separators).
            # Extract every quoted single char with a regex.
            for ch in re.findall(r'"."', raw_text):
                if ch not in seen_single:
                    seen_single.add(ch)
                    single.append(ch)
            continue

        # ── keywords block ─────────────────────────────────────────────────
        if text == "Names of reserved words –" or text == "Names of reserved words -":
            in_keywords_block = True
            _flush_token()
            continue

        # ── named token header ─────────────────────────────────────────────
        m = _TOKEN_HDR_RE.match(text)
        if m:
            _flush_token()
            in_single_char_block = False
            token_name = _clean(m.group(1))
            pending_token = {"name": token_name, "desc": ""}
            continue

        # ── description for pending token (first Normal paragraph) ─────────
        if pending_token and not pending_token["desc"] and sname in ("", "Normal"):
            # Strip leading clause numbers like "12.2.1\t..."
            desc = re.sub(r"^\d[\d.]*\s+", "", text)
            pending_token["desc"] = desc

    _flush_token()
    return {"named": named, "single": single, "keywords": keywords}


def _format_tokens(result: dict) -> str:  # noqa: F401  (kept for programmatic use)
    lines = []

    if result["named"]:
        lines.append("-- Named lexical items")
        lines.append("")
        for t in result["named"]:
            lines.append(t["name"])
            if t["desc"]:
                lines.append(f"    {t['desc']}")
            lines.append("")

    if result["single"]:
        lines.append("-- Single-character tokens")
        lines.append("    " + "  ".join(result["single"]))
        lines.append("")

    if result["keywords"]:
        kws = sorted(result["keywords"])
        lines.append("-- Reserved words (keywords)")
        # Wrap at ~80 chars
        row, row_len = [], 0
        rows = []
        for kw in kws:
            if row_len + len(kw) + 1 > 76:
                rows.append("    " + " ".join(row))
                row, row_len = [], 0
            row.append(kw)
            row_len += len(kw) + 1
        if row:
            rows.append("    " + " ".join(row))
        lines.extend(rows)
        lines.append("")

    return "\n".join(lines)

# ── docx → Markdown conversion ─────────────────────────────────────────────────

# Paragraph styles rendered as fenced code blocks (consecutive runs are merged
# into a single block), keyed by nothing — membership only, order doesn't matter.
_CODE_STYLES = {
    "BNF", "BNF Continue", "ASN.1", "ASN.1 Continue", "ASN.1 Cont.",
    "ASN.1 Italic", "Equation",
}
_QUOTE_STYLES = {"Note", "Note 1", "Note 2", "Note 3"}
_LIST_STYLES = {"enumlev1": 0, "enumlev2": 1, "sgmLI1": 0}
_HEADING_STYLES = {
    "Heading 1": 1, "Heading 2": 2, "Heading 2A": 2, "Heading 3": 3,
    "Heading": 1, "sgmH0": 1, "Annex_Title": 1,
}
_EMPHASIS_STYLES = {
    "Table_Title", "Table_NoTitle", "Figure_Title",
    "Annex_Ref", "Annex_ref", "Rec_ISO_#", "Rec_CCITT_#",
}


def _table_to_markdown(tbl_element, doc) -> str:
    """Render a docx table element as a GitHub-flavored Markdown table."""
    table = Table(tbl_element, doc)
    rows = [
        [_clean(cell.text).replace("\n", "<br>") or " " for cell in row.cells]
        for row in table.rows
    ]
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [" "] * (width - len(r)) for r in rows]

    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def convert_to_markdown(docx_path: str) -> str:
    """Convert a spec .docx to Markdown, walking paragraphs and tables in
    document order. Heading styles become '#'-headings, BNF/ASN.1-styled
    paragraphs are merged into fenced code blocks, Note-styled paragraphs
    become blockquotes, enumlev1/2 become (nested) list items, and tables
    become pipe tables. Anything else is emitted as a plain paragraph.

    Front matter (title page, Summary, Source, table of contents, ...) is
    dropped: output starts at the "# Introduction" heading. If no such
    heading is found, the full document is returned.
    """
    doc = Document(docx_path)
    para_by_elem = {p._p: p for p in doc.paragraphs}

    out: list[str] = []
    code_buf: list[str] = []

    def _flush_code():
        if code_buf:
            out.append("```asn1\n" + "\n".join(code_buf) + "\n```")
            code_buf.clear()

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "tbl":
            _flush_code()
            md = _table_to_markdown(child, doc)
            if md:
                out.append(md)
            continue

        if tag != "p":
            continue

        p = para_by_elem.get(child)
        if p is None:
            continue

        text = _clean(p.text)
        if not text:
            continue

        style = p.style.name if p.style else ""

        if style in _CODE_STYLES:
            code_buf.append(text)
            continue
        _flush_code()

        if style in _HEADING_STYLES:
            out.append("#" * _HEADING_STYLES[style] + " " + text)
        elif style in _QUOTE_STYLES:
            out.append("> " + text)
        elif style in _LIST_STYLES:
            out.append("  " * _LIST_STYLES[style] + "- " + text)
        elif style in _EMPHASIS_STYLES:
            out.append(f"**{text}**")
        else:
            out.append(text)

    _flush_code()

    try:
        out = out[out.index("# Introduction"):]
    except ValueError:
        pass  # no Introduction heading found; keep the full document

    return "\n\n".join(out) + "\n"

# ── main ──────────────────────────────────────────────────────────────────────

def main() -> int:
    args = sys.argv[1:]

    if "--clean" in args:
        clean()
        return 0

    docx_files = [a for a in args if a.endswith(".docx")]
    if not docx_files:
        # No explicit files supplied — run the full download + convert pipeline.
        docx_files = download_all()
    if not docx_files:
        print("No .docx files found.", file=sys.stderr)
        return 1

    # Grammar
    with open(GRAMMAR_FILE, "w", encoding="utf-8") as out:
        for path in docx_files:
            name = os.path.basename(path)
            prods = extract_grammar(path)
            header = f"-- {name} ({len(prods)} productions)"
            print(header)
            out.write(f"{header}\n\n")
            for p in prods:
                out.write(_format_prod(p) + "\n\n")
    print(f"Written {GRAMMAR_FILE}")

    # Tokens
    with open(TOKENS_FILE, "w", encoding="utf-8") as out:
        all_keywords: list[str] = []
        all_single: list[str] = []
        seen_single: set[str] = set()

        for path in docx_files:
            name = os.path.basename(path)
            result = extract_tokens(path)

            n_named = len(result["named"])
            n_single = len(result["single"])
            n_kw = len(result["keywords"])
            header = f"-- {name} ({n_named} named, {n_single} single-char, {n_kw} keywords)"
            print(header)
            out.write(f"{header}\n\n")

            # Named items always per-file (de-dup within file already)
            if result["named"]:
                out.write("-- Named lexical items\n\n")
                for t in result["named"]:
                    out.write(t["name"] + "\n")
                    if t["desc"]:
                        out.write(f"    {t['desc']}\n")
                    out.write("\n")

            # Accumulate single-char and keywords globally (they're only in X6801)
            for s in result["single"]:
                if s not in seen_single:
                    seen_single.add(s)
                    all_single.append(s)
            all_keywords.extend(result["keywords"])

        if all_single:
            out.write("-- Single-character tokens\n")
            out.write("    " + "  ".join(all_single) + "\n\n")

        if all_keywords:
            kws = sorted(set(all_keywords))
            out.write("-- Reserved words (keywords)\n")
            row, row_len = [], 0
            for kw in kws:
                if row_len + len(kw) + 1 > 76:
                    out.write("    " + " ".join(row) + "\n")
                    row, row_len = [], 0
                row.append(kw)
                row_len += len(kw) + 1
            if row:
                out.write("    " + " ".join(row) + "\n")
            out.write("\n")

    print(f"Written {TOKENS_FILE}")

    # Markdown (one file per spec, next to its .docx)
    for path in docx_files:
        md_path = os.path.splitext(path)[0] + ".md"
        markdown = convert_to_markdown(path)
        with open(md_path, "w", encoding="utf-8") as out:
            out.write(markdown)
        print(f"Written {md_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
